import streamlit as st
from groq import Groq
import json
import time
import hashlib
import csv
import io
from datetime import datetime, timedelta
import PyPDF2
import os
import sqlite3
import logging
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Load environment variables
load_dotenv()

# Setup logging
logging.basicConfig(
    filename="resume_analyser.log",
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

# Get API key from environment
API_KEY = os.getenv("GROQ_API_KEY")

if not API_KEY or API_KEY == "gsk_your_api_key_here_replace_with_actual_key":
    st.error("❌ API key not configured. See README.md for setup instructions.")
    st.stop()

# ── Database Setup ────────────────────────────────────────────────────────────
def init_database():
    """Initialize SQLite database for storing analyses."""
    conn = sqlite3.connect("resume_analyses.db")
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS analyses (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
            jd_hash TEXT,
            resume_hash TEXT,
            score INTEGER,
            verdict TEXT,
            summary TEXT,
            matched_skills TEXT,
            missing_skills TEXT,
            strengths TEXT,
            weaknesses TEXT,
            rewrite_tips TEXT,
            interview_questions TEXT,
            model_used TEXT,
            full_result TEXT
        )
    """)
    conn.commit()
    conn.close()

init_database()

def save_analysis_to_db(jd: str, resume: str, result: dict, model_used: str):
    """Save analysis result to database."""
    try:
        conn = sqlite3.connect("resume_analyses.db")
        cursor = conn.cursor()
        jd_hash = hashlib.md5(jd.encode()).hexdigest()
        resume_hash = hashlib.md5(resume.encode()).hexdigest()
        
        cursor.execute("""
            INSERT INTO analyses 
            (jd_hash, resume_hash, score, verdict, summary, matched_skills, missing_skills, 
             strengths, weaknesses, rewrite_tips, interview_questions, model_used, full_result)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            jd_hash, resume_hash, result.get("score"), result.get("verdict"), 
            result.get("summary"), json.dumps(result.get("matched_skills", [])),
            json.dumps(result.get("missing_skills", [])), json.dumps(result.get("strengths", [])),
            json.dumps(result.get("weaknesses", [])), json.dumps(result.get("rewrite_tips", [])),
            json.dumps(result.get("interview_questions", [])), model_used, json.dumps(result)
        ))
        conn.commit()
        conn.close()
        logging.info(f"Analysis saved - Score: {result.get('score')}, Model: {model_used}")
    except Exception as e:
        logging.error(f"Error saving to database: {str(e)}")

def get_analysis_history(limit=10):
    """Retrieve recent analyses from database."""
    try:
        conn = sqlite3.connect("resume_analyses.db")
        cursor = conn.cursor()
        cursor.execute("""
            SELECT timestamp, score, verdict, model_used FROM analyses 
            ORDER BY timestamp DESC LIMIT ?
        """, (limit,))
        results = cursor.fetchall()
        conn.close()
        return results
    except Exception as e:
        logging.error(f"Error retrieving history: {str(e)}")
        return []

# ── Rate Limiting ──────────────────────────────────────────────────────────────
def check_rate_limit(limit_seconds=30):
    """Check if user can make another analysis request."""
    if "last_analysis_time" not in st.session_state:
        st.session_state.last_analysis_time = None
    
    current_time = datetime.now()
    if st.session_state.last_analysis_time:
        elapsed = (current_time - st.session_state.last_analysis_time).total_seconds()
        if elapsed < limit_seconds:
            remaining = limit_seconds - int(elapsed)
            return False, remaining
    
    st.session_state.last_analysis_time = current_time
    return True, 0

# ── Input Validation ────────────────────────────────────────────────────────────
def validate_input(jd: str, resume: str) -> tuple:
    """Validate JD and resume inputs."""
    errors = []
    
    jd = jd.strip()
    resume = resume.strip()
    
    if len(jd) < 50:
        errors.append("Job Description too short (min 50 characters)")
    if len(resume) < 50:
        errors.append("Resume too short (min 50 characters)")
    if len(jd) > 10000:
        errors.append("Job Description too long (max 10,000 characters)")
    if len(resume) > 10000:
        errors.append("Resume too long (max 10,000 characters)")
    
    return jd, resume, errors
st.set_page_config(
    page_title="AI Resume Analyser",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ── Initialize session state ──────────────────────────────────────────────────
if "analysis_history" not in st.session_state:
    st.session_state.analysis_history = []
if "cache" not in st.session_state:
    st.session_state.cache = {}

# ── Custom CSS (Dark Mode) ─────────────────────────────────────────────────────
st.markdown("""
<style>
    body {
        background-color: #1a1a1a;
        color: #ffffff;
    }
    
    .main { 
        padding: 1rem 2rem; 
        background-color: #1a1a1a;
        color: #ffffff;
    }
    
    .score-box {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 1.5rem;
        border-radius: 12px;
        text-align: center;
        color: white;
    }
    .score-number {
        font-size: 3.5rem;
        font-weight: 800;
        line-height: 1;
    }
    .score-label {
        font-size: 0.9rem;
        opacity: 0.85;
        margin-top: 4px;
    }
    .skill-tag {
        display: inline-block;
        background: #3d3d3d;
        color: #e0e0e0;
        padding: 3px 12px;
        border-radius: 20px;
        font-size: 0.82rem;
        margin: 3px 3px;
        font-weight: 500;
    }
    .skill-tag-match {
        background: #1a4d2e;
        color: #90ee90;
    }
    .section-header {
        font-size: 1rem;
        font-weight: 600;
        color: #ffffff;
        margin-bottom: 0.5rem;
        padding-bottom: 4px;
        border-bottom: 2px solid #444444;
    }
    .tip-box {
        background: #1e3a4a;
        border-left: 4px solid #3b82f6;
        padding: 0.75rem 1rem;
        border-radius: 0 8px 8px 0;
        margin: 6px 0;
        font-size: 0.88rem;
        color: #87ceeb;
    }
    .verdict-strong { color: #90ee90; font-weight: 700; font-size: 1.05rem; }
    .verdict-medium { color: #ffa500; font-weight: 700; font-size: 1.05rem; }
    .verdict-weak   { color: #ff6b6b; font-weight: 700; font-size: 1.05rem; }
    div[data-testid="stTextArea"] textarea {
        font-size: 0.85rem;
        background-color: #2d2d2d;
        color: #ffffff;
        border-color: #444444;
    }
    .history-item {
        background: #2d2d2d;
        padding: 0.75rem 1rem;
        border-radius: 6px;
        margin: 0.5rem 0;
        font-size: 0.85rem;
        border-left: 3px solid #3b82f6;
        color: #ffffff;
    }
    .example-box {
        background: #2d2d2d;
        border: 1px solid #444444;
        padding: 0.75rem 1rem;
        border-radius: 6px;
        font-size: 0.85rem;
        color: #ffffff;
    }
</style>
""", unsafe_allow_html=True)


# ── Helpers ───────────────────────────────────────────────────────────────────
def get_client(api_key: str) -> Groq:
    return Groq(api_key=api_key)


def generate_cache_key(jd: str, resume: str) -> str:
    """Generate a unique key for caching results."""
    combined = f"{jd}||{resume}"
    return hashlib.md5(combined.encode()).hexdigest()


def extract_text_from_pdf(uploaded_file) -> str:
    """Extract text from a PDF file."""
    try:
        pdf_reader = PyPDF2.PdfReader(uploaded_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text.strip()
    except Exception as e:
        st.error(f"❌ Error reading PDF: {str(e)}")
        return ""


def analyse(client: Groq, jd: str, resume: str, model: str = "mixtral-8x7b-32768", temperature: float = 0.3, max_tokens: int = 5000, retries: int = 3) -> dict:
    """Analyse resume with comprehensive, multi-dimensional analysis."""
    prompt = f"""You are a senior HR expert and recruiter. Return ONLY valid JSON (no markdown, no explanation):
{{
  "score": 8,
  "verdict": "Good Match",
  "summary": "Brief assessment",
  "matched_skills": ["skill1", "skill2"],
  "missing_skills": ["skill1", "skill2"],
  "skill_match_percentage": 75,
  "strengths": ["str1", "str2", "str3"],
  "weaknesses": ["weak1", "weak2"],
  "experience_level": "Mid-Level",
  "experience_alignment": "well-aligned",
  "ats_compatibility": "compatible",
  "cultural_fit": "good fit",
  "remote_readiness": "good",
  "technical_depth": "intermediate",
  "soft_skills_assessment": "good",
  "education_match": "meets requirements",
  "certification_recommendations": ["cert1", "cert2"],
  "career_progression": "steady growth",
  "keyword_optimization": "good keyword coverage",
  "value_proposition": "competitive",
  "market_demand": "in-demand",
  "salary_expectation": "competitive",
  "risk_factors": ["risk1"],
  "hidden_opportunities": ["opp1"],
  "rewrite_tips": ["tip1", "tip2", "tip3", "tip4", "tip5", "tip6"],
  "quick_wins": ["win1", "win2"],
  "interview_questions": ["q1?", "q2?", "q3?", "q4?", "q5?", "q6?", "q7?", "q8?"]
}}

Now analyze this:

JOB DESCRIPTION: {jd}

RESUME: {resume}

Provide detailed analysis for each field above. Keep answers concise but meaningful."""
    
    for attempt in range(retries):
        try:
            res = client.chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": prompt}],
                temperature=temperature,
                max_tokens=max_tokens,
            )
            raw = res.choices[0].message.content.strip()
            
            # Strip markdown code fences
            if raw.startswith("```"):
                raw = raw.split("```")[1]
                if raw.startswith("json"):
                    raw = raw[4:]
            if raw.endswith("```"):
                raw = raw[:-3]
            
            # Parse JSON
            result = json.loads(raw.strip())
            
            # Ensure all required fields exist
            defaults = {
                "score": 5,
                "verdict": "Average Match",
                "summary": "Analysis complete",
                "matched_skills": [],
                "missing_skills": [],
                "skill_match_percentage": 50,
                "strengths": [],
                "weaknesses": [],
                "experience_level": "Unknown",
                "experience_alignment": "unknown",
                "ats_compatibility": "compatible",
                "cultural_fit": "unknown fit",
                "remote_readiness": "moderate",
                "technical_depth": "intermediate",
                "soft_skills_assessment": "moderate",
                "education_match": "meets requirements",
                "certification_recommendations": [],
                "career_progression": "steady growth",
                "keyword_optimization": "good keyword coverage",
                "value_proposition": "competitive",
                "market_demand": "in-demand",
                "salary_expectation": "competitive",
                "risk_factors": [],
                "hidden_opportunities": [],
                "rewrite_tips": [],
                "quick_wins": [],
                "interview_questions": []
            }
            
            # Fill missing fields with defaults
            for key, default_val in defaults.items():
                if key not in result:
                    result[key] = default_val
            
            return result
            
        except json.JSONDecodeError as e:
            if attempt < retries - 1:
                wait_time = 2 ** attempt
                st.warning(f"⚠️ Attempt {attempt + 1} failed (JSON parse error). Retrying in {wait_time}s...")
                logging.error(f"JSON parsing failed on attempt {attempt + 1}: {str(e)[:100]}")
                time.sleep(wait_time)
            else:
                st.error(f"❌ Failed after {retries} attempts. JSON parsing error.")
                logging.error(f"Final JSON parse error: {str(e)}")
                raise
                
        except Exception as e:
            if attempt < retries - 1:
                wait_time = 2 ** attempt
                error_type = type(e).__name__
                st.warning(f"⚠️ Attempt {attempt + 1} failed ({error_type}). Retrying in {wait_time}s...")
                logging.error(f"Error on attempt {attempt + 1}: {str(e)[:100]}")
                time.sleep(wait_time)
            else:
                error_msg = str(e)[:200]
                st.error(f"❌ Analysis failed: {error_msg}")
                logging.error(f"Final error: {error_msg}")
                raise


def score_color(score: int) -> str:
    if score >= 8:
        return "#16a34a"
    elif score >= 6:
        return "#d97706"
    else:
        return "#dc2626"


def verdict_class(verdict: str) -> str:
    v = verdict.lower()
    if "strong" in v:
        return "verdict-strong"
    elif "good" in v:
        return "verdict-strong"
    elif "average" in v:
        return "verdict-medium"
    else:
        return "verdict-weak"


def export_to_csv(result: dict, jd: str, resume: str) -> str:
    """Export comprehensive analysis results to CSV format."""
    output = io.StringIO()
    writer = csv.writer(output)
    
    writer.writerow(["COMPREHENSIVE RESUME ANALYSIS REPORT"])
    writer.writerow([f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"])
    writer.writerow([])
    
    # Score Section
    writer.writerow(["OVERALL SCORE & VERDICT"])
    writer.writerow(["Match Score", result.get("score", "N/A")])
    writer.writerow(["Verdict", result.get("verdict", "N/A")])
    writer.writerow(["Summary", result.get("summary", "N/A")])
    writer.writerow([])
    
    # Quick Metrics
    writer.writerow(["QUICK METRICS"])
    writer.writerow(["Skill Match %", result.get("skill_match_percentage", "N/A")])
    writer.writerow(["Experience Alignment", result.get("experience_alignment", "N/A")])
    writer.writerow(["ATS Compatibility", result.get("ats_compatibility", "N/A")])
    writer.writerow(["Market Demand", result.get("market_demand", "N/A")])
    writer.writerow(["Keyword Optimization", result.get("keyword_optimization", "N/A")])
    writer.writerow([])
    
    # Skills
    writer.writerow(["MATCHED SKILLS"])
    for skill in result.get("matched_skills", []):
        writer.writerow([skill])
    writer.writerow([])
    
    writer.writerow(["MISSING SKILLS"])
    for skill in result.get("missing_skills", []):
        writer.writerow([skill])
    writer.writerow([])
    
    # Strengths & Weaknesses
    writer.writerow(["STRENGTHS"])
    for strength in result.get("strengths", []):
        writer.writerow([strength])
    writer.writerow([])
    
    writer.writerow(["WEAKNESSES"])
    for weakness in result.get("weaknesses", []):
        writer.writerow([weakness])
    writer.writerow([])
    
    # Career Analysis
    writer.writerow(["CAREER ANALYSIS"])
    writer.writerow(["Experience Level", result.get("experience_level", "N/A")])
    writer.writerow(["Education Match", result.get("education_match", "N/A")])
    writer.writerow(["Technical Depth", result.get("technical_depth", "N/A")])
    writer.writerow(["Soft Skills", result.get("soft_skills_assessment", "N/A")])
    writer.writerow(["Cultural Fit", result.get("cultural_fit", "N/A")])
    writer.writerow(["Remote Readiness", result.get("remote_readiness", "N/A")])
    writer.writerow(["Career Progression", result.get("career_progression", "N/A")])
    writer.writerow(["Salary Expectation", result.get("salary_expectation", "N/A")])
    writer.writerow([])
    
    # Value Proposition
    writer.writerow(["VALUE PROPOSITION"])
    writer.writerow([result.get("value_proposition", "N/A")])
    writer.writerow([])
    
    # Quick Wins
    quick_wins = result.get("quick_wins", [])
    if quick_wins:
        writer.writerow(["QUICK WINS (High-Impact, Easy Changes)"])
        for win in quick_wins:
            writer.writerow([win])
        writer.writerow([])
    
    # Risk Factors
    risks = result.get("risk_factors", [])
    if risks:
        writer.writerow(["RISK FACTORS"])
        for risk in risks:
            writer.writerow([risk])
        writer.writerow([])
    
    # Hidden Opportunities
    opps = result.get("hidden_opportunities", [])
    if opps:
        writer.writerow(["HIDDEN OPPORTUNITIES"])
        for opp in opps:
            writer.writerow([opp])
        writer.writerow([])
    
    # Certifications
    certs = result.get("certification_recommendations", [])
    if certs:
        writer.writerow(["CERTIFICATION RECOMMENDATIONS"])
        for cert in certs:
            writer.writerow([cert])
        writer.writerow([])
    
    # Improvement Tips
    writer.writerow(["RESUME IMPROVEMENT TIPS"])
    for i, tip in enumerate(result.get("rewrite_tips", []), 1):
        writer.writerow([f"Tip {i}", tip])
    writer.writerow([])
    
    # Interview Questions
    writer.writerow(["INTERVIEW PREPARATION QUESTIONS"])
    for i, question in enumerate(result.get("interview_questions", []), 1):
        writer.writerow([f"Q{i}", question])
    
    return output.getvalue()


def export_to_docx(result: dict) -> bytes:
    """Export analysis results to DOCX format with comprehensive analysis."""
    doc = Document()
    
    # Title
    title = doc.add_heading('📄 Comprehensive Resume Analysis Report', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Date
    date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M')}")
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()  # Blank line
    
    # Score section
    doc.add_heading('📊 Match Score', level=1)
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Light Grid Accent 1'
    
    table.rows[0].cells[0].text = "Overall Score"
    table.rows[0].cells[1].text = f"{result.get('score', 'N/A')}/10"
    table.rows[1].cells[0].text = "Verdict"
    table.rows[1].cells[1].text = result.get('verdict', 'N/A')
    table.rows[2].cells[0].text = "Summary"
    table.rows[2].cells[1].text = result.get('summary', 'N/A')
    
    doc.add_paragraph()
    
    # Quick Metrics
    doc.add_heading('⚡ Quick Match Metrics', level=2)
    metrics_table = doc.add_table(rows=5, cols=2)
    metrics_table.style = 'Light Grid Accent 1'
    metrics_table.rows[0].cells[0].text = "Skill Match %"
    metrics_table.rows[0].cells[1].text = f"{result.get('skill_match_percentage', 0)}%"
    metrics_table.rows[1].cells[0].text = "Experience Alignment"
    metrics_table.rows[1].cells[1].text = result.get('experience_alignment', 'Unknown')
    metrics_table.rows[2].cells[0].text = "ATS Compatibility"
    metrics_table.rows[2].cells[1].text = result.get('ats_compatibility', 'Unknown')
    metrics_table.rows[3].cells[0].text = "Market Demand"
    metrics_table.rows[3].cells[1].text = result.get('market_demand', 'Unknown')
    metrics_table.rows[4].cells[0].text = "Keyword Optimization"
    metrics_table.rows[4].cells[1].text = result.get('keyword_optimization', 'Unknown')
    
    doc.add_paragraph()
    
    # Matched Skills
    doc.add_heading('✅ Matched Skills', level=2)
    for skill in result.get('matched_skills', []):
        p = doc.add_paragraph(skill, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
    
    # Missing Skills
    doc.add_heading('❌ Missing Skills', level=2)
    for skill in result.get('missing_skills', []):
        p = doc.add_paragraph(skill, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
    
    # Strengths
    doc.add_heading('💪 Strengths', level=2)
    for strength in result.get('strengths', []):
        p = doc.add_paragraph(strength, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
    
    # Weaknesses
    doc.add_heading('⚠️ Areas to Improve', level=2)
    for weakness in result.get('weaknesses', []):
        p = doc.add_paragraph(weakness, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
    
    # Career Analysis
    doc.add_heading('📈 Career Analysis', level=2)
    career_table = doc.add_table(rows=8, cols=2)
    career_table.style = 'Light Grid Accent 1'
    career_table.rows[0].cells[0].text = "Experience Level"
    career_table.rows[0].cells[1].text = result.get('experience_level', 'Unknown')
    career_table.rows[1].cells[0].text = "Education Match"
    career_table.rows[1].cells[1].text = result.get('education_match', 'Unknown')
    career_table.rows[2].cells[0].text = "Technical Depth"
    career_table.rows[2].cells[1].text = result.get('technical_depth', 'Unknown')
    career_table.rows[3].cells[0].text = "Soft Skills"
    career_table.rows[3].cells[1].text = result.get('soft_skills_assessment', 'Unknown')
    career_table.rows[4].cells[0].text = "Cultural Fit"
    career_table.rows[4].cells[1].text = result.get('cultural_fit', 'Unknown')
    career_table.rows[5].cells[0].text = "Remote Readiness"
    career_table.rows[5].cells[1].text = result.get('remote_readiness', 'Unknown')
    career_table.rows[6].cells[0].text = "Career Progression"
    career_table.rows[6].cells[1].text = result.get('career_progression', 'Unknown')
    career_table.rows[7].cells[0].text = "Salary Expectation"
    career_table.rows[7].cells[1].text = result.get('salary_expectation', 'Unknown')
    
    doc.add_paragraph()
    
    # Value Proposition
    doc.add_heading('🎯 Your Unique Value Proposition', level=2)
    doc.add_paragraph(result.get('value_proposition', 'Unable to determine'))
    
    # Quick Wins
    quick_wins = result.get('quick_wins', [])
    if quick_wins:
        doc.add_heading('⚡ Quick Wins', level=2)
        for win in quick_wins:
            p = doc.add_paragraph(win, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
    
    # Risk Factors
    risks = result.get('risk_factors', [])
    if risks:
        doc.add_heading('🚨 Risk Factors to Address', level=2)
        for risk in risks:
            p = doc.add_paragraph(risk, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
    
    # Hidden Opportunities
    opps = result.get('hidden_opportunities', [])
    if opps:
        doc.add_heading('💡 Hidden Opportunities', level=2)
        for opp in opps:
            p = doc.add_paragraph(opp, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
    
    # Certifications
    certs = result.get('certification_recommendations', [])
    if certs:
        doc.add_heading('📚 Certification Recommendations', level=2)
        for cert in certs:
            p = doc.add_paragraph(cert, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
    
    # Improvement Tips
    doc.add_heading('✏️ Resume Improvement Tips', level=2)
    for i, tip in enumerate(result.get('rewrite_tips', []), 1):
        p = doc.add_paragraph(f"{tip}", style='List Number')
        p.paragraph_format.left_indent = Inches(0.25)
    
    # Interview Questions
    doc.add_heading('🎤 Interview Preparation Questions', level=2)
    for i, question in enumerate(result.get('interview_questions', []), 1):
        p = doc.add_paragraph(f"{question}", style='List Number')
        p.paragraph_format.left_indent = Inches(0.25)
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()



# ── UI ────────────────────────────────────────────────────────────────────────
st.title("📄 AI Resume Analyser")
st.caption("Powered by Groq API · Select your preferred model in the sidebar")

# ── SIDEBAR ────────────────────────────────────────────────────────────────────
with st.sidebar:
    st.header("⚙️ Configuration")
    
    # API Key status
    st.success("✅ API Key Loaded from .env file")
    st.info("""
    **🔐 API Key Security:**
    
    Your API key is safely stored in the `.env` file (not in code).
    
    **Never commit .env to GitHub!**
    It's already in `.gitignore` for your safety.
    """)
    
    st.markdown("---")
    
    st.markdown("---")
    
    # AI Configuration
    st.subheader("⚙️ AI Model Settings")
    
    # Model selection
    available_models = [
        "llama-3.1-8b-instant",
        "llama-3.3-70b-versatile",
        "openai/gpt-oss-20b",
        "openai/gpt-oss-120b",
    ]
    selected_model = st.selectbox(
        "Select Model",
        available_models,
        index=0,
        help="Choose from available Groq production models. Fastest: llama-3.1-8b-instant, Most powerful: llama-3.3-70b-versatile"
    )
    
    temperature = st.slider(
        "Temperature (creativity)",
        min_value=0.0,
        max_value=1.0,
        value=0.3,
        step=0.1,
        help="Lower = more focused, Higher = more creative"
    )
    max_tokens = st.slider(
        "Max tokens (response length)",
        min_value=500,
        max_value=2000,
        value=1000,
        step=100,
        help="Maximum length of AI response"
    )
    
    st.markdown("---")
    
    # Analysis history from database
    st.markdown("---")
    col_refresh, col_clear = st.columns(2)
    with col_refresh:
        if st.button("🔄 Refresh History"):
            st.rerun()
    with col_clear:
        if st.button("🗑️ Clear All History"):
            try:
                conn = sqlite3.connect("resume_analyses.db")
                cursor = conn.cursor()
                cursor.execute("DELETE FROM analyses")
                conn.commit()
                conn.close()
                st.success("✅ History cleared!")
                st.rerun()
            except Exception as e:
                st.error(f"❌ Error clearing history: {str(e)}")
    
    history = get_analysis_history(limit=5)
    if history:
        st.subheader("📊 Recent Analyses")
        for timestamp, score, verdict, model in history:
            st.markdown(f'<div class="history-item">Score: {score}/10 · {verdict} · Model: {model.split("/")[-1][:12]} · {timestamp}</div>', unsafe_allow_html=True)
    else:
        st.write("No analyses yet. Start by analyzing a resume!")
    
    st.markdown("---")
    
    # Model info
    st.subheader("📊 Model Info")
    col1, col2 = st.columns(2)
    with col1:
        st.metric("Selected", selected_model.split("/")[-1][:15], "✅")
    with col2:
        st.metric("Cost", "FREE", "⚡")
    
    st.info("💡 **Model Guide:**\n\n"
            "• **llama-3.1-8b-instant** (⚡ Fastest): Good for quick analysis\n"
            "• **llama-3.3-70b-versatile** (🧠 Most powerful): Best quality analysis\n"
            "• **openai/gpt-oss-20b**: Fast & balanced\n"
            "• **openai/gpt-oss-120b**: Most powerful (slower)")
    
    st.markdown("---")


# ── INPUT COLUMNS ──────────────────────────────────────────────────────────────
col1, col2 = st.columns(2)

# SAMPLE DATA
SAMPLE_JD = """Senior Data Analyst - Remote

Company: Tech Startup Inc.

About the role:
We're looking for an experienced Data Analyst to join our growing team. You'll work with stakeholders to transform raw data into actionable insights.

Requirements:
- 3+ years experience with SQL and Python
- Advanced Excel and data visualization skills
- Experience with Tableau or Power BI
- Strong communication and presentation skills
- Knowledge of statistical analysis
- Experience with cloud platforms (AWS/GCP)

Preferred:
- Experience with machine learning concepts
- Familiarity with Git/version control
- Agile methodology experience"""

SAMPLE_RESUME = """Sai
Email: sai@email.com | LinkedIn: linkedin.com/in/sai | GitHub: github.com/sai-analyst

EDUCATION
B.Tech in Computer Science - Delhi University (2022)
GPA: 3.8/4.0

SKILLS
Languages: Python, SQL, R
Tools: Excel (Advanced), Tableau, Git, Jupyter Notebook
Cloud: AWS (EC2, S3), Basic GCP
Libraries: Pandas, NumPy, Matplotlib, Scikit-learn

EXPERIENCE

Data Analyst Intern - Analytics Co. (Jan 2022 - Jul 2022)
- Analyzed 5+ datasets using SQL and Python
- Created interactive Tableau dashboards for 10+ stakeholders
- Improved data processing efficiency by 30% with Python scripts
- Presented weekly insights to management team

PROJECTS
1. Sales Dashboard (Tableau + Excel)
   - Built interactive dashboard tracking 100+ sales metrics
   - Identified trends leading to 15% revenue increase

2. Customer Segmentation (Python + Pandas)
   - Clustered customers using K-means algorithm
   - Delivered actionable insights for marketing team

3. Data Pipeline Automation (Python + SQL)
   - Automated daily data imports reducing manual work by 20 hours/week
   - Built data validation checks preventing errors

CERTIFICATIONS
- Google Analytics Professional Certificate (2022)
- AWS Cloud Practitioner (In Progress)"""

with col1:
    st.markdown("### 📋 Job Description")
    jd_tab1, jd_tab2 = st.tabs(["📝 Paste Text", "📄 Upload PDF"])
    
    with jd_tab1:
        if st.button("📌 Load Example JD", key="load_jd"):
            st.session_state.example_jd_loaded = True
        
        jd_text = st.text_area(
            "Paste the job description here",
            height=320,
            placeholder="Example: Senior Data Analyst...",
            value=SAMPLE_JD if st.session_state.get("example_jd_loaded", False) else "",
            label_visibility="collapsed"
        )
        
        if jd_text and st.session_state.get("example_jd_loaded", False):
            st.success("✅ Example loaded!")
    
    with jd_tab2:
        jd_file = st.file_uploader("Upload JD PDF", type="pdf", key="jd_pdf")
        jd_from_pdf = ""
        if jd_file:
            jd_from_pdf = extract_text_from_pdf(jd_file)
            if jd_from_pdf:
                st.success(f"✅ PDF uploaded! ({len(jd_from_pdf)} characters)")
                st.text_area("Extracted JD text:", value=jd_from_pdf, height=200, disabled=True)
    
    jd = jd_from_pdf if jd_from_pdf else jd_text

with col2:
    st.markdown("### 👤 Your Resume")
    resume_tab1, resume_tab2 = st.tabs(["📝 Paste Text", "📄 Upload PDF"])
    
    with resume_tab1:
        if st.button("📌 Load Example Resume", key="load_resume"):
            st.session_state.example_resume_loaded = True
        
        resume_text = st.text_area(
            "Paste your resume text here",
            height=320,
            placeholder="Example: Name: ...",
            value=SAMPLE_RESUME if st.session_state.get("example_resume_loaded", False) else "",
            label_visibility="collapsed"
        )
        
        if resume_text and st.session_state.get("example_resume_loaded", False):
            st.success("✅ Example loaded!")
    
    with resume_tab2:
        resume_file = st.file_uploader("Upload Resume PDF", type="pdf", key="resume_pdf")
        resume_from_pdf = ""
        if resume_file:
            resume_from_pdf = extract_text_from_pdf(resume_file)
            if resume_from_pdf:
                st.success(f"✅ PDF uploaded! ({len(resume_from_pdf)} characters)")
                st.text_area("Extracted Resume text:", value=resume_from_pdf, height=200, disabled=True)
    
    resume = resume_from_pdf if resume_from_pdf else resume_text

# Analyse button
analyse_btn = st.button(
    "🔍 Analyse Resume",
    use_container_width=True,
    type="primary",
    disabled=not (jd and resume)
)

if not (jd and resume):
    st.info("👈 Fill in both job description and resume to get started.", icon="ℹ️")

# ── Results ────────────────────────────────────────────────────────────────────
if analyse_btn:
    # Validate inputs
    jd_clean, resume_clean, errors = validate_input(jd, resume)
    
    if errors:
        for error in errors:
            st.error(f"❌ {error}")
    else:
        # Check rate limit
        can_proceed, remaining = check_rate_limit(limit_seconds=30)
        
        if not can_proceed:
            st.warning(f"⏳ Please wait {remaining}s before next analysis (rate limit)")
        else:
            with st.spinner("🔄 Analysing your resume with AI (with automatic retries)... ⚡"):
                try:
                    client = get_client(API_KEY)
                    result = analyse(client, jd_clean, resume_clean, model=selected_model, temperature=temperature, max_tokens=max_tokens, retries=3)
                    
                    # Save to database
                    save_analysis_to_db(jd_clean, resume_clean, result, selected_model)
                    
                    # Cache result
                    cache_key = generate_cache_key(jd_clean, resume_clean)
                    st.session_state.cache[cache_key] = result
                    time.sleep(0.3)
                except json.JSONDecodeError:
                    st.error("❌ AI returned an unexpected format. This might be due to API limits. Please try again in a moment.")
                    st.stop()
                except Exception as e:
                    error_msg = str(e).lower()
                    st.error(f"❌ Error: {str(e)}")
                    st.info("💡 **Troubleshooting Tips:**\n\n"
                           "1. **Check API Key:** Verify your .env file has valid `GROQ_API_KEY`\n"
                           "2. **Try Different Model:** Select another model from dropdown\n"
                           "3. **API Limits:** Free tier has rate limits\n"
                           "4. **Wait and Retry:** Try again in a few moments")
                    logging.error(f"Analysis error: {str(e)}")
                    st.stop()

        st.markdown("---")
        st.markdown("## 📊 Analysis Results")

        # Top row — score + verdict + summary
        top1, top2 = st.columns([1, 3])

        with top1:
            color = score_color(result.get("score", 5))
            st.markdown(f"""
            <div class="score-box" style="background: linear-gradient(135deg, {color}cc, {color});">
                <div class="score-number">{result.get("score", "–")}/10</div>
                <div class="score-label">Match Score</div>
            </div>
            """, unsafe_allow_html=True)

        with top2:
            vc = verdict_class(result.get("verdict", ""))
            st.markdown(f'<p class="{vc}">{result.get("verdict", "")}</p>', unsafe_allow_html=True)
            st.markdown(result.get("summary", ""))

        st.markdown("---")

        # 📊 QUICK METRICS
        st.markdown('<p class="section-header">📊 Quick Match Metrics</p>', unsafe_allow_html=True)
        col_m1, col_m2, col_m3, col_m4 = st.columns(4)
        
        with col_m1:
            skill_pct = result.get("skill_match_percentage", 0)
            st.metric("Skill Match", f"{skill_pct}%", "📈")
        
        with col_m2:
            exp_align = result.get("experience_alignment", "unknown")
            st.metric("Experience", exp_align, "💼")
        
        with col_m3:
            ats_score = result.get("ats_compatibility", "unknown")
            st.metric("ATS Compatible", ats_score, "🤖")
        
        with col_m4:
            market = result.get("market_demand", "unknown")
            st.metric("Market Demand", market, "📈")
        
        st.markdown("---")

        # ✅ MATCHED SKILLS
        st.markdown('<p class="section-header">✅ Skills You Already Have</p>', unsafe_allow_html=True)
        matched = result.get("matched_skills", [])
        if matched:
            tags = " ".join([f'<span class="skill-tag skill-tag-match">{s}</span>' for s in matched])
            st.markdown(tags, unsafe_allow_html=True)
        else:
            st.write("No clear skill matches found.")
        
        st.markdown("---")

        # ❌ MISSING SKILLS
        st.markdown('<p class="section-header">❌ Missing Skills to Learn</p>', unsafe_allow_html=True)
        missing = result.get("missing_skills", [])
        if missing:
            tags = " ".join([f'<span class="skill-tag">{s}</span>' for s in missing])
            st.markdown(tags, unsafe_allow_html=True)
        else:
            st.write("No major missing skills!")
        
        st.markdown("---")

        # 💪 STRENGTHS & WEAKNESSES
        sw1, sw2 = st.columns(2)

        with sw1:
            st.markdown('<p class="section-header">💪 Your Strengths</p>', unsafe_allow_html=True)
            for s in result.get("strengths", []):
                st.markdown(f"✔ {s}")

        with sw2:
            st.markdown('<p class="section-header">⚠️ Areas to Improve</p>', unsafe_allow_html=True)
            for w in result.get("weaknesses", []):
                st.markdown(f"• {w}")

        st.markdown("---")

        # 📈 DETAILED ANALYSIS
        st.markdown('<p class="section-header">📈 Detailed Career Analysis</p>', unsafe_allow_html=True)
        
        col_d1, col_d2 = st.columns(2)
        
        with col_d1:
            st.write("**📊 Experience Level**")
            st.info(result.get("experience_level", "Unknown"))
        
        with col_d2:
            st.write("**🎓 Education Match**")
            st.info(result.get("education_match", "Unknown"))
        
        col_d3, col_d4 = st.columns(2)
        
        with col_d3:
            st.write("**🧠 Technical Depth**")
            st.info(result.get("technical_depth", "Unknown"))
        
        with col_d4:
            st.write("**🤝 Soft Skills**")
            st.info(result.get("soft_skills_assessment", "Unknown"))
        
        col_d5, col_d6 = st.columns(2)
        
        with col_d5:
            st.write("**🏢 Cultural Fit**")
            st.info(result.get("cultural_fit", "Unknown"))
        
        with col_d6:
            st.write("**💻 Remote Readiness**")
            st.info(result.get("remote_readiness", "Unknown"))
        
        col_d7, col_d8 = st.columns(2)
        
        with col_d7:
            st.write("**📈 Career Progression**")
            st.info(result.get("career_progression", "Unknown"))
        
        with col_d8:
            st.write("**💰 Salary Range**")
            st.info(result.get("salary_expectation", "Unknown"))
        
        st.markdown("---")

        # ⚡ QUICK WINS
        quick_wins = result.get("quick_wins", [])
        if quick_wins:
            st.markdown('<p class="section-header">⚡ Quick Wins (Easy High-Impact Changes)</p>', unsafe_allow_html=True)
            for win in quick_wins:
                st.success(f"✨ {win}")
            st.markdown("---")

        # 🔑 KEYWORD OPTIMIZATION
        st.markdown('<p class="section-header">🔑 Keyword & ATS Optimization</p>', unsafe_allow_html=True)
        st.info(f"**ATS Compatibility:** {result.get('ats_compatibility', 'Unknown')}")
        st.info(f"**Keyword Optimization:** {result.get('keyword_optimization', 'Unknown')}")
        st.markdown("---")

        # 🎯 VALUE PROPOSITION
        st.markdown('<p class="section-header">🎯 Your Unique Value Proposition</p>', unsafe_allow_html=True)
        st.warning(result.get("value_proposition", "Unable to determine"))
        st.markdown("---")

        # 🚨 RISK FACTORS
        risk_factors = result.get("risk_factors", [])
        if risk_factors:
            st.markdown('<p class="section-header">🚨 Potential Risk Factors (Address These!)</p>', unsafe_allow_html=True)
            for risk in risk_factors:
                st.error(f"⚠️ {risk}")
            st.markdown("---")

        # 💡 HIDDEN OPPORTUNITIES
        opportunities = result.get("hidden_opportunities", [])
        if opportunities:
            st.markdown('<p class="section-header">💡 Hidden Opportunities to Leverage</p>', unsafe_allow_html=True)
            for opp in opportunities:
                st.success(f"🚀 {opp}")
            st.markdown("---")

        # 📚 CERTIFICATION RECOMMENDATIONS
        certs = result.get("certification_recommendations", [])
        if certs:
            st.markdown('<p class="section-header">📚 Certification Recommendations</p>', unsafe_allow_html=True)
            for cert in certs:
                st.info(f"📜 {cert}")
            st.markdown("---")

        # ✏️ RESUME IMPROVEMENT TIPS
        st.markdown('<p class="section-header">✏️ How to Improve Your Resume (All Tips)</p>', unsafe_allow_html=True)
        for tip in result.get("rewrite_tips", []):
            st.markdown(f'<div class="tip-box">💡 {tip}</div>', unsafe_allow_html=True)

        st.markdown("---")

        # 🎤 INTERVIEW QUESTIONS
        st.markdown('<p class="section-header">🎤 Interview Preparation Questions</p>', unsafe_allow_html=True)
        for i, q in enumerate(result.get("interview_questions", []), 1):
            st.markdown(f"**Q{i}.** {q}")

        st.markdown("---")
        
        # Download and share section
        col_dl1, col_dl2, col_dl3 = st.columns(3)
        
        with col_dl1:
            csv_data = export_to_csv(result, jd_clean, resume_clean)
            st.download_button(
                label="📥 Download CSV",
                data=csv_data,
                file_name=f"resume_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv",
                use_container_width=True
            )
        
        with col_dl2:
            docx_data = export_to_docx(result)
            st.download_button(
                label="📄 Download DOCX",
                data=docx_data,
                file_name=f"resume_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        
        with col_dl3:
            st.button(
                "📋 Copy Results",
                use_container_width=True,
                help="Copy analysis to clipboard"
            )
        
        st.markdown("---")
        st.success("✅ Analysis complete! Download your results above.")
        st.caption("Built with Groq API · A fresher project by Kamparapu Eswar Sai Kiran")
